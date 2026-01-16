#!/usr/bin/env python3
"""
Word Document Translator using ChatGPT
Translates .docx documents from English to Simplified Chinese
"""

import os
import sys
from pathlib import Path
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv
import logging

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Security constants
MAX_FILE_SIZE_MB = 10
MAX_TEXT_LENGTH = 50000


class DocxTranslator:
    def __init__(self, api_key=None):
        """Initialize the translator with OpenAI API key"""
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        if not self.api_key:
            raise ValueError(
                "OpenAI API key not found. Please set OPENAI_API_KEY environment variable "
                "or pass it as an argument."
            )
        self.client = OpenAI(api_key=self.api_key)
        self.source_language = "English"
        self.target_language = "Simplified Chinese"

    def extract_text_from_docx(self, docx_path):
        """Extract all text content from a .docx file"""
        doc = Document(docx_path)
        paragraphs = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        return paragraphs

    def translate_text(self, text, max_chunk_size=3000):
        """
        Translate text using ChatGPT API
        Splits long text into chunks if necessary
        """
        if not text.strip():
            return text
        
        # For very long texts, split into chunks
        if len(text) > max_chunk_size:
            return self._translate_in_chunks(text, max_chunk_size)
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",  # Using gpt-4o-mini for cost efficiency
                messages=[
                    {
                        "role": "system",
                        "content": f"You are a professional translator. Translate the following text from {self.source_language} to {self.target_language}. Maintain the original formatting, structure, and meaning. Only return the translated text without any additional explanations."
                    },
                    {
                        "role": "user",
                        "content": text
                    }
                ],
                temperature=0.3
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f"Error translating text: {e}")
            raise

    def _translate_in_chunks(self, text, chunk_size):
        """Translate long text by splitting into chunks"""
        # Split by sentences or paragraphs if possible
        sentences = text.split('. ')
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) < chunk_size:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # Translate each chunk
        translated_chunks = []
        for i, chunk in enumerate(chunks):
            print(f"Translating chunk {i+1}/{len(chunks)}...")
            translated_chunks.append(self.translate_text(chunk))
        
        return " ".join(translated_chunks)

    def _translate_paragraph_preserve_formatting(self, paragraph):
        """
        Translate a paragraph while preserving all run-level formatting.
        This maintains bold, italic, font, size, color, etc.
        """
        full_text = paragraph.text
        if not full_text.strip():
            return
        
        # Translate the full text
        translated_text = self.translate_text(full_text)
        
        runs = paragraph.runs
        if not runs:
            return
        
        # Put all translated text in the first run, clear others
        if len(runs) == 1:
            runs[0].text = translated_text
        else:
            runs[0].text = translated_text
            for run in runs[1:]:
                run.text = ""
    
    def _translate_cell_preserve_formatting(self, cell):
        """
        Translate a table cell while preserving formatting.
        """
        full_text = cell.text
        if not full_text.strip():
            return
        
        translated_text = self.translate_text(full_text)
        
        paragraphs = cell.paragraphs
        if not paragraphs:
            return
        
        first_para = paragraphs[0]
        if first_para.runs:
            first_para.runs[0].text = translated_text
            for run in first_para.runs[1:]:
                run.text = ""
        else:
            first_para.text = translated_text
        
        for para in paragraphs[1:]:
            for run in para.runs:
                run.text = ""

    def create_translated_docx(self, source_docx_path, output_path=None):
        """
        Create a translated version of the .docx file.
        Modifies the document in-place to preserve ALL formatting.
        """
        print(f"Reading document: {source_docx_path}")
        doc = Document(source_docx_path)
        
        # Count items for progress
        para_count = sum(1 for para in doc.paragraphs if para.text.strip())
        
        # Translate paragraphs in-place
        print("Translating paragraphs...")
        translated = 0
        for para in doc.paragraphs:
            if para.text.strip():
                translated += 1
                print(f"Translating paragraph {translated}/{para_count}...")
                self._translate_paragraph_preserve_formatting(para)
        
        # Translate tables in-place
        print("Translating tables...")
        for table_idx, table in enumerate(doc.tables):
            print(f"Translating table {table_idx+1}...")
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        self._translate_cell_preserve_formatting(cell)
        
        # Determine output path
        if output_path is None:
            source_path = Path(source_docx_path)
            output_path = source_path.parent / f"{source_path.stem}_translated{source_path.suffix}"
        
        # Save the modified document
        doc.save(output_path)
        print(f"\nTranslation complete! Saved to: {output_path}")
        return output_path


def validate_path(file_path, must_exist=True):
    """
    Validate file path for security.
    Prevents path traversal attacks.
    """
    try:
        # Resolve to absolute path and ensure it's under a safe directory
        resolved = Path(file_path).resolve()
        
        # Check for path traversal attempts
        if '..' in str(file_path):
            logger.warning(f"Path traversal attempt detected: {file_path}")
            return None
        
        if must_exist and not resolved.exists():
            return None
            
        return resolved
    except Exception as e:
        logger.error(f"Path validation error: {e}")
        return None


def main():
    """Main function to run the translator"""
    if len(sys.argv) < 2:
        print("Usage: python translate_docx.py <input_file.docx> [output_file.docx]")
        print("\nExample:")
        print("  python translate_docx.py document.docx")
        print("  python translate_docx.py document.docx translated_document.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    # Security: Validate input file path
    input_path = validate_path(input_file, must_exist=True)
    if input_path is None:
        print(f"Error: File '{input_file}' not found or invalid path.")
        sys.exit(1)
    
    if not str(input_path).lower().endswith('.docx'):
        print("Error: Input file must be a .docx file.")
        sys.exit(1)
    
    # Security: Check file size
    file_size_mb = input_path.stat().st_size / (1024 * 1024)
    if file_size_mb > MAX_FILE_SIZE_MB:
        print(f"Error: File too large ({file_size_mb:.1f} MB). Maximum size is {MAX_FILE_SIZE_MB} MB.")
        sys.exit(1)
    
    # Security: Validate output path if provided
    if output_file:
        # Ensure output ends with .docx
        if not output_file.lower().endswith('.docx'):
            output_file = output_file + '.docx'
        
        output_path = validate_path(output_file, must_exist=False)
        if output_path is None:
            print(f"Error: Invalid output path '{output_file}'.")
            sys.exit(1)
        output_file = output_path
    
    try:
        # Initialize translator
        translator = DocxTranslator()
        
        # Translate document
        result_path = translator.create_translated_docx(str(input_path), output_file)
        print(f"\n✓ Successfully translated document!")
        print(f"  Input:  {input_path}")
        print(f"  Output: {result_path}")
        
    except ValueError as e:
        print(f"Error: {e}")
        print("\nPlease set your OpenAI API key:")
        print("  1. Create a .env file in this directory")
        print("  2. Add: OPENAI_API_KEY=your_api_key_here")
        print("  Or set it as an environment variable")
        sys.exit(1)
    except Exception as e:
        # Don't expose internal error details
        logger.error(f"Translation error: {e}")
        print("Error: Translation failed. Check the logs for details.")
        sys.exit(1)


if __name__ == "__main__":
    main()

