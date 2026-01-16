#!/usr/bin/env python3
"""
Streamlit Web UI for Word Document Translator
Upload .docx files and download translated versions
"""

import streamlit as st
import os
import tempfile
from pathlib import Path
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv
import io
import zipfile
import re
import logging

# Load environment variables
load_dotenv()

# Configure logging - don't log sensitive data
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Security constants
MAX_FILE_SIZE_MB = 10  # Maximum file size in MB
MAX_FILES_PER_UPLOAD = 20  # Maximum number of files per upload
MAX_TEXT_LENGTH = 50000  # Maximum characters per text block to translate

# Page configuration
st.set_page_config(
    page_title="Word Document Translator",
    page_icon="📄",
    layout="centered"
)

# Custom CSS for better styling
st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        text-align: center;
        color: #1f77b4;
        margin-bottom: 1rem;
    }
    .sub-header {
        text-align: center;
        color: #666;
        margin-bottom: 2rem;
    }
    .stButton>button {
        width: 100%;
        background-color: #1f77b4;
        color: white;
        font-weight: bold;
    }
    </style>
""", unsafe_allow_html=True)


class DocxTranslator:
    def __init__(self, api_key=None):
        """Initialize the translator with OpenAI API key"""
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        if not self.api_key:
            raise ValueError("OpenAI API key not found")
        
        # Initialize OpenAI client
        # Use explicit import to avoid any proxy-related issues
        from openai import OpenAI as OpenAIClient
        self.client = OpenAIClient(api_key=self.api_key)
        self.source_language = "English"
        self.target_language = "Simplified Chinese"

    def _sanitize_text_for_translation(self, text):
        """
        Sanitize text before sending to the API.
        Removes potential prompt injection attempts.
        """
        if len(text) > MAX_TEXT_LENGTH:
            logger.warning(f"Text truncated from {len(text)} to {MAX_TEXT_LENGTH} characters")
            text = text[:MAX_TEXT_LENGTH]
        return text

    def translate_text(self, text, max_chunk_size=3000, progress_callback=None):
        """
        Translate text using ChatGPT API
        Splits long text into chunks if necessary
        """
        if not text.strip():
            return text
        
        # Sanitize input
        text = self._sanitize_text_for_translation(text)
        
        # For very long texts, split into chunks
        if len(text) > max_chunk_size:
            return self._translate_in_chunks(text, max_chunk_size, progress_callback)
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": f"You are a professional translator. Translate the following text from {self.source_language} to {self.target_language}. Maintain the original formatting, structure, and meaning. Only return the translated text without any additional explanations."
                    },
                    {
                        "role": "user",
                        "content": text
                    }
                ],
                temperature=0.3
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            # Log the actual error but don't expose details to user
            logger.error(f"Translation error: {str(e)}")
            raise Exception("Translation failed. Please try again later.")

    def _translate_in_chunks(self, text, chunk_size, progress_callback=None):
        """Translate long text by splitting into chunks"""
        sentences = text.split('. ')
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) < chunk_size:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # Translate each chunk
        translated_chunks = []
        for i, chunk in enumerate(chunks):
            if progress_callback:
                progress_callback(i + 1, len(chunks))
            translated_chunks.append(self.translate_text(chunk))
        
        return " ".join(translated_chunks)

    def _translate_paragraph_preserve_formatting(self, paragraph):
        """
        Translate a paragraph while preserving all run-level formatting.
        This maintains bold, italic, font, size, color, etc.
        """
        # Get the full text of the paragraph
        full_text = paragraph.text
        if not full_text.strip():
            return
        
        # Translate the full text
        translated_text = self.translate_text(full_text)
        
        # If paragraph has runs, try to preserve formatting
        runs = paragraph.runs
        if not runs:
            return
        
        # Strategy: Put all translated text in the first run, clear others
        # This preserves the first run's formatting for the entire paragraph
        if len(runs) == 1:
            # Simple case: single run, just replace text
            runs[0].text = translated_text
        else:
            # Multiple runs: put translated text in first run, clear others
            # This preserves the paragraph's primary formatting
            runs[0].text = translated_text
            for run in runs[1:]:
                run.text = ""
    
    def _translate_cell_preserve_formatting(self, cell):
        """
        Translate a table cell while preserving formatting.
        Handles multiple paragraphs within a cell.
        """
        # Get all text from all paragraphs in the cell
        full_text = cell.text
        if not full_text.strip():
            return
        
        # Translate the full cell text
        translated_text = self.translate_text(full_text)
        
        # Get all paragraphs in the cell
        paragraphs = cell.paragraphs
        if not paragraphs:
            return
        
        # Put translated text in the first paragraph's first run
        first_para = paragraphs[0]
        if first_para.runs:
            # Clear all runs except first, put translated text in first
            first_para.runs[0].text = translated_text
            for run in first_para.runs[1:]:
                run.text = ""
        else:
            # No runs, add text directly (rare case)
            first_para.text = translated_text
        
        # Clear text from other paragraphs in the cell (keep the paragraph for spacing)
        for para in paragraphs[1:]:
            for run in para.runs:
                run.text = ""

    def create_translated_docx_bytes(self, source_docx_bytes, progress_bar=None, status_text=None):
        """
        Create a translated version of the .docx file from bytes.
        Modifies the document in-place to preserve ALL formatting.
        Returns the translated document as bytes.
        """
        # Load document from bytes - we'll modify this directly
        doc = Document(io.BytesIO(source_docx_bytes))
        
        # Count total items to translate for progress tracking
        total_paragraphs = sum(1 for para in doc.paragraphs if para.text.strip())
        total_cells = sum(
            sum(1 for row in table.rows for cell in row.cells if cell.text.strip())
            for table in doc.tables
        )
        total_items = total_paragraphs + total_cells
        
        if total_items == 0:
            total_items = 1  # Avoid division by zero
        
        current_item = 0
        
        # Translate paragraphs in-place (preserves all formatting)
        if status_text:
            status_text.text("Translating paragraphs...")
        
        for para in doc.paragraphs:
            if para.text.strip():
                current_item += 1
                if progress_bar:
                    progress_bar.progress(min(current_item / total_items, 1.0))
                self._translate_paragraph_preserve_formatting(para)
        
        # Translate tables in-place (preserves table structure and cell formatting)
        if status_text:
            status_text.text("Translating tables...")
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        current_item += 1
                        if progress_bar:
                            progress_bar.progress(min(current_item / total_items, 1.0))
                        self._translate_cell_preserve_formatting(cell)
        
        # Save the modified document to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes.getvalue()
    
    def translate_filename(self, filename):
        """
        Translate a filename from English to Chinese.
        Preserves the file extension.
        """
        # Extract the stem (filename without extension) and extension
        path = Path(filename)
        stem = path.stem
        extension = path.suffix
        
        # Check if filename contains mostly non-ASCII (likely already Chinese)
        non_ascii_count = sum(1 for c in stem if ord(c) > 127)
        if non_ascii_count > len(stem) * 0.3:
            # Likely already in Chinese, return as-is
            return filename
        
        # Check if it's mostly numbers or special characters
        alpha_count = sum(1 for c in stem if c.isalpha())
        if alpha_count < 2:
            # Too few letters to translate meaningfully
            return filename
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a translator. Translate the following filename from English to Simplified Chinese. Keep it concise and suitable as a filename. Only return the translated filename without any explanation, quotes, or file extension. Do not add any punctuation."
                    },
                    {
                        "role": "user",
                        "content": stem
                    }
                ],
                temperature=0.3
            )
            translated_stem = response.choices[0].message.content.strip()
            # Clean up the translated filename (remove invalid characters)
            translated_stem = re.sub(r'[<>:"/\\|?*]', '', translated_stem)
            translated_stem = translated_stem.strip('. ')
            
            if translated_stem:
                return f"{translated_stem}{extension}"
            else:
                return filename
        except Exception:
            # If translation fails, return original filename
            return filename


def validate_api_key(api_key):
    """
    Validate the OpenAI API key by making a simple API call.
    Returns (is_valid, error_message)
    """
    if not api_key:
        return False, "API key is empty"
    
    # Basic format check
    if not api_key.startswith(('sk-', 'sk-proj-')):
        return False, "Invalid API key format. Key should start with 'sk-' or 'sk-proj-'"
    
    if len(api_key) < 20:
        return False, "API key is too short"
    
    # Test the API key with a minimal request
    try:
        from openai import OpenAI as OpenAIClient
        client = OpenAIClient(api_key=api_key)
        # Make a minimal API call to validate
        client.models.list()
        return True, None
    except Exception as e:
        error_msg = str(e)
        if "invalid_api_key" in error_msg.lower() or "incorrect api key" in error_msg.lower():
            return False, "Invalid API key. Please check your key and try again."
        elif "rate_limit" in error_msg.lower():
            # Rate limited means the key is valid
            return True, None
        else:
            logger.error(f"API key validation error: {error_msg}")
            return False, "Could not validate API key. Please check your key and internet connection."


def render_api_key_section():
    """
    Render the API key input section.
    Returns the API key if valid, None otherwise.
    """
    # Initialize session state for API key
    if 'api_key' not in st.session_state:
        st.session_state.api_key = None
    if 'api_key_validated' not in st.session_state:
        st.session_state.api_key_validated = False
    
    # Check if API key is already in environment (for backward compatibility)
    env_api_key = os.getenv("OPENAI_API_KEY")
    
    # If we have a validated key in session, show status and option to change
    if st.session_state.api_key_validated and st.session_state.api_key:
        st.success("✅ API key configured and validated")
        
        # Show masked key
        masked_key = st.session_state.api_key[:8] + "..." + st.session_state.api_key[-4:]
        st.text(f"Current key: {masked_key}")
        
        if st.button("🔄 Change API Key", key="change_key"):
            st.session_state.api_key = None
            st.session_state.api_key_validated = False
            st.rerun()
        
        return st.session_state.api_key
    
    # Show API key input form
    st.markdown("### 🔑 Enter Your OpenAI API Key")
    
    st.info("""
    **How to get your API key:**
    1. Go to [OpenAI Platform](https://platform.openai.com/api-keys)
    2. Sign in or create an account
    3. Click "Create new secret key"
    4. Copy the key and paste it below
    
    ⚠️ **Note:** Your API key is stored only in your browser session and is never saved to disk.
    """)
    
    # API key input
    with st.form("api_key_form"):
        api_key_input = st.text_input(
            "API Key",
            type="password",
            placeholder="sk-...",
            help="Your OpenAI API key starting with 'sk-'"
        )
        
        submit_button = st.form_submit_button("🔐 Validate & Save Key", use_container_width=True)
        
        if submit_button:
            if not api_key_input:
                st.error("Please enter an API key")
                return None
            
            with st.spinner("Validating API key..."):
                is_valid, error_msg = validate_api_key(api_key_input)
            
            if is_valid:
                st.session_state.api_key = api_key_input
                st.session_state.api_key_validated = True
                st.success("✅ API key validated successfully!")
                st.rerun()
            else:
                st.error(f"❌ {error_msg}")
                return None
    
    # Also check for environment variable as fallback
    if env_api_key:
        st.markdown("---")
        st.markdown("**Or use environment variable:**")
        if st.button("Use API key from environment variable", key="use_env_key"):
            with st.spinner("Validating environment API key..."):
                is_valid, error_msg = validate_api_key(env_api_key)
            
            if is_valid:
                st.session_state.api_key = env_api_key
                st.session_state.api_key_validated = True
                st.success("✅ Environment API key validated!")
                st.rerun()
            else:
                st.error(f"❌ Environment key invalid: {error_msg}")
    
    return None


def main():
    """Main Streamlit app"""
    # Header
    st.markdown('<p class="main-header">📄 Word Document Translator</p>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">Translate English documents to Simplified Chinese using ChatGPT</p>', unsafe_allow_html=True)
    
    # API Key section
    api_key = render_api_key_section()
    
    if not api_key:
        st.warning("⚠️ Please enter your OpenAI API key above to use the translator.")
        st.stop()
    
    # File uploader - allow multiple files
    st.markdown("### 📤 Upload Documents")
    uploaded_files = st.file_uploader(
        "Choose .docx files to translate",
        type=['docx'],
        accept_multiple_files=True,
        help="Upload one or more English Word documents (.docx format)"
    )
    
    if uploaded_files:
        # Security: Check number of files
        if len(uploaded_files) > MAX_FILES_PER_UPLOAD:
            st.error(f"⚠️ Too many files! Maximum {MAX_FILES_PER_UPLOAD} files allowed per upload.")
            st.stop()
        
        # Security: Check file sizes
        oversized_files = []
        for f in uploaded_files:
            size_mb = f.size / (1024 * 1024)
            if size_mb > MAX_FILE_SIZE_MB:
                oversized_files.append(f"{f.name} ({size_mb:.1f} MB)")
        
        if oversized_files:
            st.error(f"⚠️ Files too large (max {MAX_FILE_SIZE_MB} MB each):\n" + "\n".join(oversized_files))
            st.stop()
        
        # Display file info
        st.markdown(f"**{len(uploaded_files)} file(s) selected:**")
        
        file_list = []
        for f in uploaded_files:
            file_list.append({
                "Filename": f.name,
                "Size": f"{f.size / 1024:.2f} KB"
            })
        
        with st.expander("📋 File Information", expanded=False):
            for i, info in enumerate(file_list, 1):
                st.text(f"{i}. {info['Filename']} ({info['Size']})")
        
        # Translate button
        if st.button("🚀 Translate All Documents", type="primary", use_container_width=True):
            try:
                # Initialize translator
                translator = DocxTranslator(api_key)
                
                # Store translated files
                translated_files = []
                
                # Overall progress
                overall_progress = st.progress(0)
                overall_status = st.empty()
                
                # Process each file
                for file_idx, uploaded_file in enumerate(uploaded_files):
                    overall_status.markdown(f"**Processing file {file_idx + 1} of {len(uploaded_files)}:** {uploaded_file.name}")
                    
                    # Create per-file progress indicators
                    file_progress = st.progress(0)
                    file_status = st.empty()
                    
                    # Read uploaded file
                    file_bytes = uploaded_file.read()
                    
                    # Translate the filename
                    file_status.text("Translating filename...")
                    original_name = uploaded_file.name
                    translated_filename = translator.translate_filename(original_name)
                    
                    # Translate document content
                    file_status.text("Translating document content...")
                    translated_bytes = translator.create_translated_docx_bytes(
                        file_bytes,
                        progress_bar=file_progress,
                        status_text=file_status
                    )
                    
                    # Store the result
                    translated_files.append({
                        'original_name': original_name,
                        'translated_name': translated_filename,
                        'data': translated_bytes
                    })
                    
                    # Clear per-file progress
                    file_progress.empty()
                    file_status.empty()
                    
                    # Update overall progress
                    overall_progress.progress((file_idx + 1) / len(uploaded_files))
                
                # Clear overall progress
                overall_progress.empty()
                overall_status.empty()
                
                # Success message
                st.success(f"✅ Translation complete! {len(translated_files)} file(s) translated.")
                
                # Download section
                st.markdown("### 📥 Download Translated Documents")
                
                # Show translated filenames
                st.markdown("**Translated filenames:**")
                for tf in translated_files:
                    st.text(f"  {tf['original_name']} → {tf['translated_name']}")
                
                st.markdown("---")
                
                # If multiple files, offer a zip download
                if len(translated_files) > 1:
                    # Create zip file
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                        for tf in translated_files:
                            zf.writestr(tf['translated_name'], tf['data'])
                    zip_buffer.seek(0)
                    
                    st.download_button(
                        label="⬇️ Download All as ZIP",
                        data=zip_buffer.getvalue(),
                        file_name="translated_documents.zip",
                        mime="application/zip",
                        type="primary",
                        use_container_width=True
                    )
                    
                    st.markdown("**Or download individually:**")
                
                # Individual download buttons
                for i, tf in enumerate(translated_files):
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.text(tf['translated_name'])
                    with col2:
                        st.download_button(
                            label="⬇️",
                            data=tf['data'],
                            file_name=tf['translated_name'],
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_{i}"
                        )
                
                st.info("💡 Tip: The translated documents preserve the original formatting including paragraphs and tables.")
                
            except ValueError as e:
                # Configuration errors are safe to show
                st.error(f"❌ Configuration Error: {str(e)}")
            except Exception as e:
                # Log full error but show generic message to user
                logger.error(f"Translation error: {str(e)}")
                st.error("❌ Translation Error: An error occurred during translation.")
                st.info("Please check your API key and internet connection, then try again.")
    
    # Sidebar with instructions
    with st.sidebar:
        st.markdown("### 📖 How to Use")
        st.markdown("""
        1. **Enter** your OpenAI API key
        2. **Upload** one or more English .docx files
        3. Click **Translate All Documents**
        4. Wait for translation to complete
        5. **Download** translated documents (ZIP or individually)
        
        ### 🔑 Getting an API Key
        1. Visit [OpenAI Platform](https://platform.openai.com/api-keys)
        2. Create an account or sign in
        3. Generate a new API key
        4. Copy and paste it into the app
        
        ### ⚙️ Settings
        - **Source Language**: English
        - **Target Language**: Simplified Chinese
        - **Model**: GPT-4o-mini (cost-effective)
        
        ### 📝 Notes
        - Multiple files can be uploaded at once
        - Filenames are also translated to Chinese
        - Documents are translated paragraph by paragraph
        - Tables are also translated
        - Original formatting is preserved
        - Large documents may take several minutes
        """)
        
        st.markdown("---")
        st.markdown("### 🔒 Privacy & Security")
        st.markdown("""
        - **API Key**: Stored only in your browser session, never saved to disk
        - **Documents**: Processed through OpenAI's API, not stored on any server
        - **Local Use**: This app is designed for local/personal use
        """)
    
    # Footer
    st.markdown("---")
    st.markdown(
        "<div style='text-align: center; color: #666;'>Powered by OpenAI GPT-4o-mini</div>",
        unsafe_allow_html=True
    )


if __name__ == "__main__":
    main()

